#!/usr/bin/env python3
"""One-off script to build quarterly_progress_keyterms_optimization.docx from the report content."""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Install: pip install python-docx")
    raise

ROOT = Path(__file__).resolve().parent

def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # Title
    t = doc.add_heading("Quarterly Progress Report: Keyterms Optimization for Speech-to-Text Biasing", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Project: ").bold = True
    p.add_run("fSTT (Forecasted STT) — retrieval-based keyword/keyterm prediction for Deepgram STT biasing")
    p = doc.add_paragraph()
    p.add_run("Reporting period: ").bold = True
    p.add_run("Q1 2025 (quarterly progress)")
    p = doc.add_paragraph()
    p.add_run("Prepared for: ").bold = True
    p.add_run("GPU cluster access renewal")
    doc.add_paragraph()

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "This quarter we advanced keyterms/keywords optimization for domain-specific speech-to-text (STT) "
        "biasing in restaurant phone-order conversations. The work focuses on predicting which terms a user "
        "is likely to say next and injecting them into Deepgram Nova-3's keyterm parameter to improve "
        "recognition of food and menu vocabulary (e.g., chicken, tonkatsu, california roll) while avoiding "
        "noise from terms that do not benefit from biasing.")
    doc.add_paragraph()

    doc.add_heading("Key Accomplishments", level=1)

    doc.add_heading("1. Keywords-focused retrieval pipeline (local restaurant + restaurant type)", level=2)
    doc.add_paragraph(
        "We implemented and trained a retrieval model (SentenceTransformers) that forecasts keywords only "
        "for the next user utterance using only local restaurant conversation data. The pipeline: uses local "
        "CSV dialogs with optional restaurant_type (e.g., asian_fusion, thai, ramen); prefixes conversation "
        "history with 'Restaurant type: X' at inference so the model conditions on cuisine/venue type; "
        "selects the best checkpoint by Recall@20 keywords; and saves a shared FAISS index plus encoder "
        "(best_model/, shared_index/) for low-latency inference in production.")
    doc.add_paragraph()

    doc.add_heading("2. Target and candidate optimization", level=2)
    doc.add_paragraph(
        "To improve quality of predicted terms: we apply aggressive stopword and easy-to-pronounce filtering "
        "(e.g., yes, 3, mm) so the candidate pool focuses on terms that benefit from biasing (e.g., dish names); "
        "we use keywords-only targets (single-word, TF-IDF unigrams plus menu vocabulary) with keyterms empty; "
        "and we exclude first-turn-only candidates from the shared index to avoid generic terms polluting retrieval.")
    doc.add_paragraph()

    doc.add_heading("3. Production integration and A/B testing", level=2)
    doc.add_paragraph(
        "We integrated the trained model into the LiveKit telephony A/B test (production/livekit_ab_test.py): "
        "Arm A uses Deepgram STT with retrieval-based keyword injection; Arm B has injection disabled. "
        "At each system turn we encode history (with Restaurant type prefix when applicable), retrieve top-k "
        "neighbors, aggregate keywords, and pass them to Deepgram Nova-3 via the keyterm parameter. We also "
        "extract and prepend explicit options from the agent's last utterance. Session transcripts and "
        "injected terms logs are saved for offline comparison.")
    doc.add_paragraph()

    doc.add_heading("4. Training and evaluation setup", level=2)
    doc.add_paragraph(
        "Training is run in Google Colab via notebooks/colab_train_retrieval_local_restaurant_type.ipynb: "
        "SentenceTransformer with MultipleNegativesRankingLoss; train/val/test split by dialog ID; "
        "validation Recall@20 keywords for checkpoint selection; artifacts saved to a timestamped run "
        "directory and used in production as model_dir for the A/B worker.")
    doc.add_paragraph()

    doc.add_heading("Technical Highlights", level=1)
    doc.add_paragraph("Model: SentenceTransformer encoder + FAISS index over candidate keyword sets.")
    doc.add_paragraph("Metric: Recall@20 keywords (validation) for checkpoint selection.")
    doc.add_paragraph("Inference: Encode history → k-NN over shared index → aggregate keywords → Deepgram keyterm.")
    doc.add_paragraph("Domain: Restaurant phone orders; optional restaurant_type conditioning.")
    doc.add_paragraph("Deployment: LiveKit agent worker; optional preload of encoder + index at startup.")
    doc.add_paragraph()

    doc.add_heading("Next Steps", level=1)
    doc.add_paragraph(
        "Scale training to larger restaurant corpora and additional restaurant types; correlate A/B transcript "
        "logs with WER/term-error metrics to quantify impact; optionally reintroduce keyterms (multi-word "
        "phrases) in a separate pipeline and compare against keywords-only.")
    doc.add_paragraph()

    doc.add_paragraph(
        "This work supports improved STT accuracy for domain-specific vocabulary in voice ordering systems "
        "and justifies continued use of GPU resources for training and experimentation.")

    out = ROOT / "quarterly_progress_keyterms_optimization.docx"
    doc.save(str(out))
    print(f"Saved: {out}")

if __name__ == "__main__":
    main()
