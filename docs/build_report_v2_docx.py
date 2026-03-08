#!/usr/bin/env python3
"""Build quarterly_progress_report_v2.docx - minimal, professional style."""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Install: pip install python-docx")
    raise

ROOT = Path(__file__).resolve().parent

def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    doc.add_paragraph("Quarterly Progress Report: Keyterms Optimization", style="Heading 1")
    doc.add_paragraph("Project: fSTT | Period: Q1 2025 | Purpose: GPU cluster access renewal")
    doc.add_paragraph()

    doc.add_paragraph("Summary", style="Heading 2")
    doc.add_paragraph(
        "We developed a retrieval-based system to predict domain-specific keywords for the next user utterance "
        "in restaurant phone-order conversations. Predicted terms are injected into Deepgram Nova-3 via the "
        "keyterm parameter to improve recognition of food and menu vocabulary. The pipeline filters out terms "
        "that do not benefit from biasing (e.g. yes, mm, numbers) and retains domain-relevant terms.")
    doc.add_paragraph()

    doc.add_paragraph("Accomplishments", style="Heading 2")
    doc.add_paragraph(
        "Retrieval pipeline. We trained a SentenceTransformer encoder to forecast keywords from conversation "
        "history. Training uses local restaurant CSV dialogs with an optional restaurant_type column. At "
        "inference, history is prefixed with \"Restaurant type: X\" so the model conditions on venue type. "
        "Best checkpoint is selected by validation Recall@20 keywords. Outputs: encoder (best_model/), FAISS "
        "index (shared_index/), candidates, metadata.")
    doc.add_paragraph(
        "Target optimization. We use TF-IDF unigrams plus optional menu vocabulary as keyword targets. "
        "Keyterms are set to empty. Aggressive filtering removes stopwords and easy-to-pronounce terms (yes, "
        "3, mm, short fillers) so the candidate pool retains terms that benefit from STT biasing. "
        "First-turn-only candidates are excluded from the shared index to avoid generic terms.")
    doc.add_paragraph(
        "Production. The model is integrated into a LiveKit telephony A/B test. Arm A runs Deepgram STT "
        "with retrieval-based keyword injection; Arm B runs without injection. On each system turn we encode "
        "history (with restaurant type prefix when applicable), retrieve top-k neighbors from the FAISS "
        "index, aggregate keywords, and pass them to Deepgram as keyterm. We also extract explicit options "
        "from the agent's last utterance (e.g. chicken, pork, shrimp) and prepend them. Transcripts and "
        "injected-terms logs are saved for offline comparison.")
    doc.add_paragraph(
        "Training. SentenceTransformer with MultipleNegativesRankingLoss; train/val/test split by dialog ID; gradient "
        "accumulation; checkpoint selection by val Recall@20 keywords. Artifacts saved to a timestamped run "
        "directory and loaded in production via model_dir.")
    doc.add_paragraph()

    doc.add_paragraph("Technical Details", style="Heading 2")
    doc.add_paragraph(
        "Model: SentenceTransformer (e.g. all-MiniLM-L6-v2) + FAISS IndexFlatIP over candidate keyword sets. "
        "Metric: Recall@20 keywords. Inference: encode history, k-NN search, aggregate keywords, cap at "
        "max_terms, pass to Deepgram keyterm. Domain: restaurant phone orders with optional restaurant_type. "
        "Deployment: LiveKit agent worker with optional encoder+index preload.")
    doc.add_paragraph()

    doc.add_paragraph("Next Steps", style="Heading 2")
    doc.add_paragraph(
        "Scale training to larger corpora and additional restaurant types. Correlate A/B logs with WER or "
        "term-error metrics. Optionally reintroduce keyterms (multi-word phrases) and compare against "
        "keywords-only. Fine-tune a custom STT model on real restaurant call data using consensus labels "
        "from two existing STT systems; the resulting model targets noisy phone environments, accented "
        "speech, and foreign menu items.")

    out = ROOT / "quarterly_progress_report_v2.docx"
    doc.save(str(out))
    print(f"Saved: {out}")

if __name__ == "__main__":
    main()
